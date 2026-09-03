from pathlib import Path
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
RELEASE = json.loads((ROOT / "functions/lib/verifactuRelease.json").read_text(encoding="utf-8"))
VERSION = RELEASE["version"]
OUT = ROOT / "output/pdf" / f"Declaracion_responsable_LimpiaGest_{VERSION}_REVISION.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
MUTED = "667085"
LIGHT = "F2F4F7"
CAUTION = "FFF4CE"
RED = "9B1C1C"


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                el = mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    mar.append(el)
                el.set(qn("w:w"), str(value))
                el.set(qn("w:type"), "dxa")


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade(cells[0], LIGHT)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, bold=True, color=DARK)
        for run in cells[1].paragraphs[0].runs:
            set_font(run)
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, text, fill=CAUTION, color=RED):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title + " ")
    set_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_font(r)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, before, after, color in (
    ("Heading 1", 16, 16, 8, BLUE),
    ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 12, 8, 4, DARK),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "LIMPIAGEST | EXPEDIENTE VERI*FACTU"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header.runs:
    set_font(run, size=8.5, color=MUTED)
footer = section.footer.paragraphs[0]
footer.text = "Borrador para revisión del productor | Producción bloqueada"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    set_font(run, size=8.5, color=MUTED, italic=True)
for alternate in (section.even_page_header, section.first_page_header):
    alternate.paragraphs[0].text = header.text
    alternate.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in alternate.paragraphs[0].runs:
        set_font(run, size=8.5, color=MUTED)
for alternate in (section.even_page_footer, section.first_page_footer):
    alternate.paragraphs[0].text = footer.text
    alternate.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in alternate.paragraphs[0].runs:
        set_font(run, size=8.5, color=MUTED, italic=True)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("DECLARACIÓN RESPONSABLE DEL SISTEMA INFORMÁTICO DE FACTURACIÓN")
set_font(r, size=22, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run(f"LimpiaGest · versión {VERSION} · BORRADOR PARA REVISIÓN")
set_font(r, size=13, color=MUTED)

add_callout(
    doc,
    "NO FIRMAR TODAVÍA.",
    "Revisión 31/08/2026. Se ha cerrado la emisión con VeriFactu desactivado y unificado la versión fiscal. Las pruebas aceptadas por la AEAT corresponden al entorno de pruebas, no acreditan por sí solas cumplimiento completo. Falta validar el alcance de la edición definitiva con remisión real antes de firmar. La producción permanece bloqueada.",
)

doc.add_heading("1. Información obligatoria", level=1)
add_kv_table(doc, [
    ("a) Nombre del sistema informático", "LimpiaGest"),
    ("b) Código identificador", "LG"),
    ("c) Versión", f"{VERSION}. Identificador común del registro fiscal y del XML de esta revisión. Edición actualmente limitada a pruebas."),
    ("d) Componentes y funcionalidades", "Aplicación web React; servicios Firebase/Google Cloud; Firestore; funciones de emisión fiscal; generación de PDF y QR; cola AEAT; conector Windows para certificados no exportables; custodia alternativa PFX/P12 mediante Secret Manager. Gestiona facturas, huellas, encadenamiento, altas, anulaciones, subsanaciones, reintentos y respuestas AEAT."),
    ("e) Funcionamiento exclusivo VERI*FACTU", "SÍ en el diseño de esta edición: el servidor impide emitir cuando VeriFactu está desactivado, tanto individualmente como por lotes. Se pueden crear y guardar borradores. La emisión disponible sigue limitada a pruebas; no se ha habilitado facturación real."),
    ("f) Uso por varios obligados tributarios", "SÍ. Arquitectura multiempresa con separación lógica y control de acceso por entidad; cada instalación/tenant mantiene su propia numeración, cadena, certificado y cola."),
    ("g) Tipos de firma de registros", "No aplicable en modalidad exclusiva VERI*FACTU. El certificado del obligado tributario se utiliza para autenticación mutua TLS en la remisión a la AEAT; los registros se protegen mediante huella SHA-256 y encadenamiento."),
    ("h) Productor", RELEASE["producer"]),
    ("i) NIF del productor", RELEASE["producerNif"]),
    ("j) Dirección postal", "C/ Algarrobo 0, 04740 Roquetas de Mar, Almería, España"),
])

doc.add_heading("2. Manifestación de cumplimiento", level=1)
doc.add_paragraph("Texto propuesto para la declaración definitiva. El productor debe verificar su exactitud y el alcance de la versión antes de suscribirlo.")
p = doc.add_paragraph()
p.add_run("k) ").bold = True
p.add_run(
    "La entidad productora manifiesta que el sistema informático LimpiaGest, en la versión indicada, cumple con lo dispuesto en el artículo 29.2.j) de la Ley 58/2003, General Tributaria; con el Reglamento aprobado por el Real Decreto 1007/2023, de 5 de diciembre; con la Orden HAC/1177/2024, de 17 de octubre; y con las especificaciones publicadas en la sede electrónica de la Agencia Estatal de Administración Tributaria que completen dicha orden."
)

doc.add_heading("3. Fecha, lugar y firma", level=1)
add_kv_table(doc, [
    ("l) Lugar", "Roquetas de Mar, España"),
    ("l) Fecha", "____ de ______________ de 20____"),
    ("Firmante", "Representante autorizado de LIMPIEZAS RAIBA SOCIEDAD LIMITADA"),
    ("Nombre y apellidos", "____________________________________________"),
    ("Firma", "\n\n____________________________________________\n"),
])

doc.add_page_break()
doc.add_heading("ANEXO I. Descripción técnica del cumplimiento", level=1)
add_kv_table(doc, [
    ("Integridad e inalterabilidad", "La emisión se ejecuta en servidor, asigna número correlativo mediante transacción, genera registro fiscal inmutable y huella SHA-256. Las facturas emitidas no se editan; las correcciones se realizan mediante los tipos fiscales previstos."),
    ("Trazabilidad", "Cada registro enlaza con el anterior y conserva identificadores de factura, empresa, versión, fecha/hora con huso, huella previa y huella resultante."),
    ("Remisión VERI*FACTU", "El registro se transforma en SOAP 1.1 oficial y se remite por TLS mutuo al endpoint de pruebas. Existe validación XSD local. El conector bloquea otros destinos. La remisión real y su puesta en servicio siguen pendientes; este borrador no las acredita."),
    ("Certificado", "El flujo local usa el certificado instalado en Windows sin exportar la clave privada. El token del conector se cifra mediante DPAPI. La alternativa cloud custodia PFX/P12 en Secret Manager."),
    ("Cola y reintentos", "Reclamación exclusiva con lease temporal, idempotencia, límite de intentos, espera exponencial y estados de aceptación, aceptación con errores, rechazo y reintento. El primer envío se bloquea si la hora fiscal lleva preparada más de tres minutos, para respetar el margen temporal de la AEAT."),
    ("Registro operativo", "Los eventos técnicos y fiscales relevantes se registran por empresa; las reglas impiden su modificación desde clientes no autorizados."),
    ("QR y leyenda", "La factura VeriFactu incorpora la URL oficial correspondiente al entorno, QR de 30 mm con corrección M, la etiqueta 'QR tributario:' y la leyenda VERI*FACTU. El caso TEST-VF-2026-0001 fue encontrado correctamente en el portal de cotejo de pruebas."),
    ("Uso multiempresa", "Los datos se segregan por companyId/tenant. Cada obligado dispone de configuración, numeración, cadena fiscal, cola, certificado y registro operativo independientes."),
])

doc.add_heading("Fuentes normativas y de consulta", level=2)
sources = [
    "Real Decreto 1007/2023, artículo 13; Orden HAC/1177/2024, artículo 15 y anexo (BOE, textos consolidados).",
    "AEAT: Certificación de los sistemas informáticos: declaración responsable. FAQ actualizadas a 21/07/2026; consultadas el 31/08/2026.",
]
for source in sources:
    p = doc.add_paragraph(source)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_font(run, size=9)

annex_evidence_heading = doc.add_heading(
    "ANEXO II. Evidencias de cierre",
    level=1,
)
annex_evidence_heading.paragraph_format.page_break_before = True
annex_evidence_heading.paragraph_format.keep_together = True
annex_evidence_heading.paragraph_format.keep_with_next = True
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ("Control", "Estado", "Evidencia / acción pendiente")
for i, value in enumerate(headers):
    table.cell(0, i).text = value
    shade(table.cell(0, i), LIGHT)
    for run in table.cell(0, i).paragraphs[0].runs:
        set_font(run, bold=True, color=DARK)
rows = [
    ("Pruebas internas", "SUPERADO", "Evidencias internas anteriores conservadas. Revisión 31/08: 36 controles focalizados de emisión, XML, bloqueo de producción y distribución del panel superados, sin envíos nuevos a la AEAT."),
    ("Certificado", "SUPERADO", "B04843843 conectado; acceso al servicio de pruebas verificado. Alternativa PFX/P12 disponible. Caducidad 20/11/2026."),
    ("Alta AEAT de pruebas", "SUPERADO", "TEST-VF-2026-0002, 1,21 EUR: aceptada sin errores al primer intento. También se conserva el alta aceptada de TEST-VF-2026-0001."),
    ("Anulación", "SUPERADO", "TEST-VF-2026-0002: aceptada sin errores al primer intento. Totales facturado y pendiente: 0,00 EUR. Se conserva el historial; no se borraron registros ni se marcaron como cobrados."),
    ("Subsanación", "SUPERADO", "TEST-VF-2026-0002: aceptada sin errores al primer intento. Nombre corregido a CLIENTE PRUEBA SUBSANACION CORREGIDO; NIF e importes sin cambios. Corrección técnica de59240 validada."),
    ("QR", "SUPERADO", "El portal AEAT de preproducción mostró 'Encontrada' y coincidieron NIF B04843843, número TEST-VF-2026-0001, fecha 31/07/2026 e importe 1,21 EUR."),
    ("Revisión independiente", "OPCIONAL", "Recomendación de apoyo al productor. La AEAT no exige certificación externa ni registro previo del software."),
    ("Emisión exclusiva", "SUPERADO", "Bloqueo en servidor antes de asignar números o escribir datos. Borradores disponibles; producción bloqueada. Versión fiscal unificada: " + VERSION + "."),
    ("Edición definitiva", "PENDIENTE", "Validar remisión real y alcance completo de la versión antes de la manifestación final. Las pruebas no sustituyen esta revisión."),
    ("Firma declaración", "PENDIENTE", "Revisar también los datos del productor y domicilio. Completar nombre, fecha y firma en la versión definitiva y conservarla accesible en la aplicación."),
    ("Instalador firmado", "PENDIENTE", "Mejora de distribución: firma de código y SmartScreen. No confundir con la declaración ni con una certificación externa obligatoria."),
]
for values in rows:
    cells = table.add_row().cells
    for i, value in enumerate(values):
        cells[i].text = value
        for run in cells[i].paragraphs[0].runs:
            pending_state = value in {"PENDIENTE", "REPETIR", "ACEPTADA CON AVISO"}
            set_font(run, size=9.5, bold=(i == 1), color=(RED if pending_state else "000000"))
set_table_geometry(table, [2350, 1550, 5460])

doc.core_properties.title = "Declaración responsable del sistema informático de facturación LimpiaGest"
doc.core_properties.subject = "Borrador de cumplimiento VERI*FACTU"
doc.core_properties.author = "LIMPIEZAS RAIBA SOCIEDAD LIMITADA"
doc.save(OUT)
print(OUT)
